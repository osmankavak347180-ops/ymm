"""Modül C — DOCX çıktı + TASLAK damgası testleri (Task 5.2).

KVKK: LLM ÇAĞRISI YOK (`uretici._llm_uret` mock). Çıktı dosyası tmp_path'e
yazılır; dosya adı `TASLAK_` önekli, her section üstbilgisinde kalın/kırmızı
"İNCELENMESİ GEREKEN TASLAK — YMM ONAYI GEREKLİDİR" damgası bulunur.
"""

from __future__ import annotations

import sqlite3
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document
from typer.testing import CliRunner

from ymm.cli import app
from ymm.db.depo import Depo
from ymm.modeller import Bulgu, Donem
from ymm.rapor.uretici import DAMGA, taslak_uret

runner = CliRunner()

_PROJE_KOKU = Path(__file__).parent.parent
_MIZAN_XLSX = _PROJE_KOKU / "ornek_veri" / "mizan_2025.xlsx"


def _kimlik_db_olustur(yol: Path) -> Path:
    baglanti = sqlite3.connect(yol)
    baglanti.execute(
        "CREATE TABLE IF NOT EXISTS kimlik "
        "(takma_kod TEXT, tip TEXT, gercek_ad TEXT, vkn_tckn TEXT)"
    )
    baglanti.execute(
        "INSERT INTO kimlik VALUES ('KISI-001', 'KISI', 'Örnek Ortak', NULL)"
    )
    baglanti.commit()
    baglanti.close()
    return yol


def _bulgu_b_131(mukellef_id: int) -> Bulgu:
    return Bulgu(
        kaynak="B",
        kontrol_kodu="B-131-ORTAK",
        seviye="yuksek",
        tutar_fark=Decimal("150000.00"),
        yuzde_fark=None,
        detay={
            "hesap_kodu": "131",
            "hesap_adi": "Ortaklardan Alacaklar [KISI-001]",
            "kural": "bakiye_var",
            "not": None,
        },
        mukellef_id=mukellef_id,
        yil=2025,
    )


@pytest.fixture
def llm_yanki(monkeypatch):
    monkeypatch.setattr(
        "ymm.rapor.uretici._llm_uret", lambda istem, sistem, kimlik_db: istem
    )


@pytest.fixture
def hazir_ortam(tmp_path, llm_yanki):
    """Dolu depo + kimlik.db + çıktı dizini."""
    depo = Depo(tmp_path / "veri.db")
    mukellef_id = depo.mukellef_ekle("MUK-001")
    depo.donem_ekle(mukellef_id, Donem(yil=2025, tip="YILLIK", sira=0))
    depo.bulgu_yaz([_bulgu_b_131(mukellef_id)])
    kimlik_db = _kimlik_db_olustur(tmp_path / "kimlik.db")
    return depo, mukellef_id, kimlik_db, tmp_path / "output"


# --- taslak_uret ---------------------------------------------------------------


def test_dosya_adi_taslak_onekli(hazir_ortam):
    depo, mukellef_id, kimlik_db, cikti = hazir_ortam

    yol = taslak_uret(
        depo, mukellef_id, 2025,
        kimlik_db=kimlik_db, takma_kod="MUK-001", cikti_dizini=cikti,
    )

    assert yol.exists()
    assert yol.name == "TASLAK_MUK-001_2025.docx"
    assert yol.name.startswith("TASLAK_")


def test_her_section_ustbilgisinde_kalin_kirmizi_damga(hazir_ortam):
    depo, mukellef_id, kimlik_db, cikti = hazir_ortam

    yol = taslak_uret(
        depo, mukellef_id, 2025,
        kimlik_db=kimlik_db, takma_kod="MUK-001", cikti_dizini=cikti,
    )

    belge = Document(str(yol))
    assert len(belge.sections) >= 1
    for section in belge.sections:
        ustbilgi_metni = "\n".join(p.text for p in section.header.paragraphs)
        assert DAMGA in ustbilgi_metni

        damga_runlari = [
            run
            for p in section.header.paragraphs
            for run in p.runs
            if DAMGA in run.text
        ]
        assert damga_runlari, "Damga tek run içinde bulunamadı"
        for run in damga_runlari:
            assert run.bold is True
            assert run.font.color.rgb is not None
            kirmizi, yesil, mavi = (
                run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2]
            )
            assert kirmizi > yesil and kirmizi > mavi  # kırmızı baskın


def test_govde_bulgu_tutarini_ve_yer_tutuculari_icerir(hazir_ortam):
    depo, mukellef_id, kimlik_db, cikti = hazir_ortam

    yol = taslak_uret(
        depo, mukellef_id, 2025,
        kimlik_db=kimlik_db, takma_kod="MUK-001", cikti_dizini=cikti,
    )

    belge = Document(str(yol))
    tam_metin = "\n".join(p.text for p in belge.paragraphs)
    assert "150.000,00 TL" in tam_metin
    assert "[YMM GÖRÜŞÜ — ELLE DOLDURULACAK]" in tam_metin
    assert "Örnek Ortak" in tam_metin  # geri-yerleştirme docx'e taşındı
    assert "[KISI-001]" not in tam_metin


# --- CLI: ymm rapor -------------------------------------------------------------


def test_cli_rapor_ucta_uca_dummy_akis(tmp_path, llm_yanki):
    """Uçtan uca: mizan yükle -> tara -> rapor -> TASLAK_ dosyası oluşur."""
    veri_db = tmp_path / "veri.db"
    kimlik_db = tmp_path / "kimlik.db"
    cikti = tmp_path / "output"

    yukle = runner.invoke(
        app,
        [
            "yukle", "mizan", str(_MIZAN_XLSX),
            "--mukellef", "MUK-001", "--yil", "2025",
            "--veri-db", str(veri_db), "--kimlik-db", str(kimlik_db),
        ],
    )
    assert yukle.exit_code == 0, yukle.output

    tara = runner.invoke(
        app,
        ["tara", "--mukellef", "MUK-001", "--yil", "2025", "--veri-db", str(veri_db)],
    )
    assert tara.exit_code == 0, tara.output

    rapor = runner.invoke(
        app,
        [
            "rapor", "--mukellef", "MUK-001", "--yil", "2025",
            "--veri-db", str(veri_db), "--kimlik-db", str(kimlik_db),
            "--cikti", str(cikti),
        ],
    )
    assert rapor.exit_code == 0, rapor.output

    dosyalar = list(cikti.glob("TASLAK_*.docx"))
    assert len(dosyalar) == 1
    assert dosyalar[0].name == "TASLAK_MUK-001_2025.docx"


def test_cli_rapor_mukellef_yoksa_hata(tmp_path):
    sonuc = runner.invoke(
        app,
        [
            "rapor", "--mukellef", "MUK-YOK", "--yil", "2025",
            "--veri-db", str(tmp_path / "veri.db"),
            "--kimlik-db", str(tmp_path / "kimlik.db"),
            "--cikti", str(tmp_path / "output"),
        ],
    )

    assert sonuc.exit_code == 1
    assert "Traceback" not in sonuc.output
