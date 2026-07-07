"""Modül C — rapor taslağı üretici, metin katmanı (Task 5.1).

Akış (docs/01-MIMARI.md §4, SKILL.md tam-tasdik-raporu):
bulguları oku -> bulgu tipi başına j2 kalıp paragrafı doldur (tutarlar Türk
biçiminde KODDA hazırlanır; LLM sayı üretmez) -> kalıpları gateway üzerinden
redakte ettir -> yanıt doğrulaması (girdideki tüm tutarlar yanıtta birebir
olmalı, değilse kalıplara güvenli geri düşüş) -> iskelet.md.j2 doldur ->
[KISI-nnn] token'larını kimlik.db'den YERELDE geri yerleştir.

KVKK:
- LLM'e giden her şey MASKELİ kalıp metnidir; gateway zaten sizinti_tara
  ile denetler. Bu modül `anthropic` IMPORT ETMEZ.
- kimlik.db'ye erişim YALNIZ `geri_yerlestir` içindedir ve LLM çağrısından
  SONRA, salt-okunur yapılır (docs/01-MIMARI.md: rapor/uretici.py bu izne
  sahip iki modülden biridir).
- Çıktı her zaman TASLAK'tır; damga metni sabittir, kaldırılamaz.
"""

from __future__ import annotations

import logging
import re
import sqlite3
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from jinja2 import Environment, FileSystemLoader

from ymm.llm.gateway import uret as _llm_uret
from ymm.llm.istemler import RAPOR_SISTEM_ISTEMI, redaksiyon_istemi
from ymm.modeller import Bulgu

_logger = logging.getLogger(__name__)

DAMGA = "İNCELENMESİ GEREKEN TASLAK — YMM ONAYI GEREKLİDİR"

_SEVIYE_SIRASI = {"yuksek": 0, "orta": 1, "dusuk": 2}

_ORTAM = Environment(
    loader=FileSystemLoader(Path(__file__).parent / "sablonlar"),
    autoescape=False,
    trim_blocks=True,
    lstrip_blocks=True,
)

# kontrol kodu -> j2 şablon dosyası (SKILL.md §2 kalıpları). Eşleşmeyen kod:
# "-ARTIS" son eki -> bulgu_b_artis.j2; aksi halde kaynak A/B genel kalıbı.
_SABLON_ESLEME = {
    "A-KDV-HASILAT": "bulgu_a_kdv_hasilat.j2",
    "A-MUHSGK-UCRET": "bulgu_a_muhsgk_ucret.j2",
    "A-GECICI-KV": "bulgu_a_gecici_kv.j2",
    "B-131-ORTAK": "bulgu_b_131_ortak.j2",
    "B-331-ORTAK": "bulgu_b_331_ortak.j2",
    "B-689-KKEG": "bulgu_b_689_kkeg.j2",
    "B-100-KASA": "bulgu_b_100_kasa.j2",
    "B-190-DEVREDEN": "bulgu_b_190_devreden.j2",
}

# Türk biçimli tutar (doğrulama için): 1.234.567,89
_TUTAR_RE = re.compile(r"\d{1,3}(?:\.\d{3})*,\d{2}")


def tutar_bicimle(deger: Decimal | str) -> str:
    """Decimal'i Türk biçimli TL string'ine çevirir: 1234567.89 ->
    '1.234.567,89 TL'. LLM'e daima bu hazır biçim verilir (SKILL.md §3)."""
    tutar = Decimal(str(deger))
    tam, ondalik = f"{tutar:.2f}".split(".")
    negatif = tam.startswith("-")
    tam = tam.lstrip("-")
    ters = tam[::-1]
    gruplu = ".".join(ters[i : i + 3] for i in range(0, len(ters), 3))[::-1]
    return f"{'-' if negatif else ''}{gruplu},{ondalik} TL"


def _yuzde_bicimle(yuzde: float | None) -> str | None:
    if yuzde is None:
        return None
    return f"{yuzde:.2f}".replace(".", ",")


def _sablon_baglami(bulgu: Bulgu) -> tuple[str, dict]:
    """Bulgu için (şablon dosyası, render bağlamı) çifti üretir. Tüm tutarlar
    burada Türk biçimine çevrilir — şablonlar/LLM sayı işlemez."""
    detay = bulgu.detay
    kod = bulgu.kontrol_kodu
    fark = tutar_bicimle(bulgu.tutar_fark) if bulgu.tutar_fark is not None else None
    yuzde = _yuzde_bicimle(bulgu.yuzde_fark)

    if kod == "A-KDV-HASILAT":
        return _SABLON_ESLEME[kod], {
            "kdv_toplam": tutar_bicimle(detay["sol_tutar"]),
            "net_satis": tutar_bicimle(detay["sag_tutar"]),
            "fark": fark,
            "yuzde": yuzde,
        }
    if kod == "A-MUHSGK-UCRET":
        return _SABLON_ESLEME[kod], {
            "muhsgk_toplam": tutar_bicimle(detay["sol_tutar"]),
            "mizan_toplam": tutar_bicimle(detay["sag_tutar"]),
            "hesap_listesi": detay.get("formul") or "ilgili gider",
            "fark": fark,
        }
    if kod == "A-GECICI-KV":
        return _SABLON_ESLEME[kod], {
            "gecici_matrah": tutar_bicimle(detay["sol_tutar"]),
            "kv_matrah": tutar_bicimle(detay["sag_tutar"]),
            "fark": fark,
        }
    if kod in _SABLON_ESLEME:  # B-özel kalıpları (131/331/689/100/190)
        return _SABLON_ESLEME[kod], {
            "hesap_kodu": detay.get("hesap_kodu", ""),
            "hesap_adi": detay.get("hesap_adi"),
            "bakiye": fark,
        }
    if kod.endswith("-ARTIS"):
        hesap_kodu = detay.get("hesap_kodu", "")
        hesap_adi = detay.get("hesap_adi")
        etiket = f"{hesap_adi} ({hesap_kodu})" if hesap_adi else hesap_kodu
        return "bulgu_b_artis.j2", {
            "hesap_etiketi": etiket,
            "onceki": tutar_bicimle(detay["onceki"]),
            "cari": tutar_bicimle(detay["cari"]),
            "yon": detay.get("yon", "artis"),
            "yuzde": yuzde,
        }
    if bulgu.kaynak == "A":
        return "bulgu_a_genel.j2", {
            "kod": kod,
            "sol": tutar_bicimle(detay["sol_tutar"]),
            "sag": tutar_bicimle(detay["sag_tutar"]),
            "fark": fark,
            "yuzde": yuzde,
            "aciklama": detay.get("aciklama"),
        }
    return "bulgu_b_genel.j2", {
        "kod": kod,
        "hesap_kodu": detay.get("hesap_kodu"),
        "hesap_adi": detay.get("hesap_adi"),
        "bakiye": fark,
        "kural_notu": detay.get("not"),
    }


def bulgu_paragraflari(bulgular: list[Bulgu]) -> list[str]:
    """Bulguları seviye sırasıyla (yüksek -> orta -> düşük) kalıp
    paragraflara döker. Her paragraf LLM'e gitmeye hazır MASKELİ metindir."""
    sirali = sorted(bulgular, key=lambda b: _SEVIYE_SIRASI.get(b.seviye, 99))
    paragraflar: list[str] = []
    for bulgu in sirali:
        sablon_adi, baglam = _sablon_baglami(bulgu)
        paragraflar.append(_ORTAM.get_template(sablon_adi).render(**baglam).strip())
    return paragraflar


def _tutarlar_korunmus(kaynak_metin: str, yanit: str) -> bool:
    """SKILL.md §4 doğrulaması: kaynak kalıplardaki TÜM Türk biçimli tutarlar
    LLM yanıtında birebir bulunmalı."""
    return all(tutar in yanit for tutar in set(_TUTAR_RE.findall(kaynak_metin)))


def _bolum_metni(bulgular: list[Bulgu], kimlik_db: Path) -> str:
    """Bir modülün (A veya B) bulgularını kalıp paragraflara döker ve gateway
    üzerinden redakte ettirir. Yanıt tutarları koruyamadıysa kalıp
    paragraflar REDAKSİYONSUZ kullanılır (güvenli geri düşüş). Bulgu yoksa
    LLM HİÇ çağrılmaz."""
    if not bulgular:
        return "Bu bölümde raporlanacak bulgu tespit edilmemiştir."

    paragraflar = bulgu_paragraflari(bulgular)
    ham = "\n\n".join(paragraflar)

    redakte = _llm_uret(redaksiyon_istemi(paragraflar), RAPOR_SISTEM_ISTEMI, kimlik_db)
    if not _tutarlar_korunmus(ham, redakte):
        _logger.warning(
            "LLM redaksiyonu girdideki tutarları birebir korumadı; kalıp "
            "paragraflara geri düşüldü (%d paragraf).",
            len(paragraflar),
        )
        return ham
    return redakte


def geri_yerlestir(metin: str, kimlik_db: Path) -> str:
    """LLM-SONRASI YEREL adım: metindeki [TAKMA-KOD] token'larını kimlik.db'deki
    gerçek adlarla değiştirir. Eşleşmeyen token AYNEN bırakılır (YMM taslakta
    takma kodu görür — sessizce silinmez)."""
    if not kimlik_db.exists():
        return metin

    baglanti = sqlite3.connect(kimlik_db)
    try:
        kayitlar = baglanti.execute(
            "SELECT takma_kod, gercek_ad FROM kimlik"
        ).fetchall()
    finally:
        baglanti.close()

    for takma_kod, gercek_ad in kayitlar:
        if takma_kod and gercek_ad:
            metin = metin.replace(f"[{takma_kod}]", gercek_ad)
    return metin


def rapor_metni_uret(
    depo, mukellef_id: int, yil: int, *, kimlik_db: Path, takma_kod: str
) -> str:
    """Rapor taslağının TAM METNİNİ (markdown) üretir — Task 5.2'de docx bu
    metinden yazılacak. Damga her zaman metnin başında ve sonunda bulunur;
    `[ELLE DOLDURULACAK]` / `[YMM GÖRÜŞÜ — ELLE DOLDURULACAK]` yer tutucuları
    korunur."""
    bulgular = depo.bulgular(mukellef_id, yil)
    a_bulgulari = [b for b in bulgular if b.kaynak == "A"]
    b_bulgulari = [b for b in bulgular if b.kaynak == "B"]

    eksik_donem_uyarilari = [
        uyari
        for bulgu in a_bulgulari
        for uyari in bulgu.detay.get("eksik_donem_uyarilari", [])
    ]

    metin = _ORTAM.get_template("iskelet.md.j2").render(
        damga=DAMGA,
        yil=yil,
        takma_kod=takma_kod,
        modul_a_metni=_bolum_metni(a_bulgulari, kimlik_db),
        modul_b_metni=_bolum_metni(b_bulgulari, kimlik_db),
        eksik_donem_uyarilari=eksik_donem_uyarilari,
    )

    return geri_yerlestir(metin, kimlik_db)


_DAMGA_RENGI = RGBColor(0xC0, 0x00, 0x00)  # koyu kırmızı


def _damga_ustbilgiye_yaz(belge: Document) -> None:
    """Her section'ın (dolayısıyla her sayfanın) üstbilgisine kalın + kırmızı
    TASLAK damgasını yazar (mimari kural 5 — kaldırılamaz)."""
    for section in belge.sections:
        paragraf = section.header.paragraphs[0]
        run = paragraf.add_run(DAMGA)
        run.bold = True
        run.font.color.rgb = _DAMGA_RENGI


def _markdown_satirini_ekle(belge: Document, satir: str) -> None:
    """rapor_metni_uret'in markdown çıktısındaki tek satırı docx'e çevirir.
    Desteklenen alt küme iskelet.md.j2'nin ürettiği kadardır — genel bir
    markdown dönüştürücü DEĞİLDİR (bilinçli sadelik)."""
    if not satir.strip() or satir.strip() == "---":
        return
    if satir.startswith("### "):
        belge.add_heading(satir[4:], level=3)
    elif satir.startswith("## "):
        belge.add_heading(satir[3:], level=2)
    elif satir.startswith("# "):
        belge.add_heading(satir[2:], level=1)
    elif satir.startswith("- "):
        belge.add_paragraph(satir[2:], style="List Bullet")
    elif satir.startswith("> "):
        run = belge.add_paragraph().add_run(satir[2:].replace("**", ""))
        run.bold = True
        run.font.color.rgb = _DAMGA_RENGI
    else:
        belge.add_paragraph(satir)


def taslak_uret(
    depo,
    mukellef_id: int,
    yil: int,
    *,
    kimlik_db: Path,
    takma_kod: str,
    cikti_dizini: Path = Path("output"),
) -> Path:
    """Bulguları okur, şablonu doldurur, LLM paragraflarını (gateway
    üzerinden) alır, kimlik geri-yerleştirir, TASLAK damgalı docx üretir
    (docs/01-MIMARI.md imzası; depo/kimlik_db/takma_kod bağımlılıkları diğer
    modüllerdeki desenle açık parametre yapıldı).

    Dosya adı her zaman ``TASLAK_{takma_kod}_{yil}.docx``; her sayfa
    üstbilgisinde kalın/kırmızı damga bulunur.
    """
    metin = rapor_metni_uret(
        depo, mukellef_id, yil, kimlik_db=kimlik_db, takma_kod=takma_kod
    )

    belge = Document()
    _damga_ustbilgiye_yaz(belge)
    for satir in metin.splitlines():
        _markdown_satirini_ekle(belge, satir)

    cikti_dizini.mkdir(parents=True, exist_ok=True)
    yol = cikti_dizini / f"TASLAK_{takma_kod}_{yil}.docx"
    belge.save(str(yol))
    return yol
